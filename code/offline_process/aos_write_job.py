#!/usr/bin/env python
# coding: utf-8

from opensearchpy import OpenSearch, RequestsHttpConnection, AWSV4SignerAuth, helpers
import boto3
import random
import json
from awsglue.utils import getResolvedOptions
import sys
import hashlib
import datetime
import re
import os
import itertools
from bs4 import BeautifulSoup
from langchain.document_loaders import PDFMinerPDFasHTMLLoader
from langchain.docstore.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter,CharacterTextSplitter, MarkdownTextSplitter
import logging
import urllib.parse
import numpy as np
from urllib.parse import unquote
from datetime import datetime

args = getResolvedOptions(sys.argv, ['bucket', 'object_key','AOS_ENDPOINT','REGION','EMB_MODEL_ENDPOINT','PUBLISH_DATE', 'company', 'emb_batch_size'])
s3 = boto3.resource('s3')
bucket = args['bucket']
object_key = args['object_key']
QA_SEP = '=====' # args['qa_sep'] # 
EXAMPLE_SEP = '\n\n'
arg_chunk_size = 1000
CHUNK_SIZE=1000
CHUNK_OVERLAP=0

EMB_MODEL_ENDPOINT=args['EMB_MODEL_ENDPOINT']
COMPANY = args.get('company', 'default')
print(f"COMPANY :{COMPANY}")
smr_client = boto3.client("sagemaker-runtime")

AOS_ENDPOINT = args['AOS_ENDPOINT']
REGION = args['REGION']

publish_date = args['PUBLISH_DATE'] if 'PUBLISH_DATE' in args.keys() else datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")

INDEX_NAME = 'chatbot-index'
EXAMPLE_INDEX_NAME = 'chatbot-example-index'
EMB_BATCH_SIZE = int(args.get('emb_batch_size', '20'))
print(f"EMB_BATCH_SIZE :{EMB_BATCH_SIZE}")
Sentence_Len_Threshold=10
Paragraph_Len_Threshold=20

DOC_INDEX_TABLE= 'chatbot_doc_index'
dynamodb = boto3.client('dynamodb')

AOS_BENCHMARK_ENABLED=False
BEDROCK_EMBEDDING_MODELID_LIST = ["cohere.embed-multilingual-v3","cohere.embed-english-v3","amazon.titan-embed-text-v1","amazon.titan-embed-text-v2:0"]


bedrock = boto3.client(service_name='bedrock-runtime',
                       region_name= os.environ.get('bedrock_region',REGION))

from anthropic_bedrock import AnthropicBedrock
anthropic_bedrock = AnthropicBedrock(
    aws_region=os.environ.get('bedrock_region',REGION),
)

def num_tokens_from_string(string: str) -> int:
    """Returns the number of tokens in a text string."""
    num_tokens = anthropic_bedrock.count_tokens(string)  # 
    # print(f"num_tokens:{num_tokens}")
    return num_tokens

def token_length_function(string:str) -> int:
    try:
        return num_tokens_from_string(string)
    except Exception as e:
        print(f'str(e),using len()')
        return len(string)

def get_embedding_from_titan(text, model_id):
    if isinstance(text, list):
        raise RuntimeError("For titan embedding, it only support str, but list[str] provided")

    body = {
        "inputText": text[:8192],  
    }

    if model_id == 'amazon.titan-embed-text-v2:0':
        body['dimensions'] = 1024 
        body['normalize'] = True

    playload = json.dumps(body)

    bedrock_resp = bedrock.invoke_model(
        body=playload,
        modelId=model_id,
        accept="application/json",
        contentType="application/json"
    )
    response_body = json.loads(bedrock_resp.get('body').read())
    return response_body['embedding']

def get_embedding_bedrock(texts,model_id):
    provider = model_id.split(".")[0]
    if provider == "cohere":
        body = json.dumps({
            "texts": [texts[:2048]] if isinstance(texts, str) else [text[:2048] for text in texts],
            "input_type": "search_document",
            # "truncate":"RIGHT" ## 该参数目前不起作用，只能通过text[:2048]来截断
        })

        bedrock_resp = bedrock.invoke_model(
            body=body,
            modelId=model_id,
            accept="application/json",
            contentType="application/json"
        )
        response_body = json.loads(bedrock_resp.get('body').read())
        embeddings = response_body['embeddings']
    elif provider == "amazon" :
        if model_id not in  ["amazon.titan-embed-text-v2:0", "amazon.titan-embed-text-v1"]:
            raise RuntimeError(f"{model_id} is not supported")

        if isinstance(texts, list) and len(texts) > 1:
            embeddings = [ get_embedding_from_titan(text, model_id) for text in texts ]
        else:
            embeddings = [ get_embedding_from_titan(texts, model_id) ]
    return embeddings


def get_embedding(smr_client, text_arrs, endpoint_name=EMB_MODEL_ENDPOINT):
    if endpoint_name in BEDROCK_EMBEDDING_MODELID_LIST:
        return get_embedding_bedrock(text_arrs,endpoint_name)
    
    if AOS_BENCHMARK_ENABLED:
        text_len = len(text_arrs)
        return [ np.random.rand(768).tolist() for i in range(text_len) ]
        
    parameters = {
    }

    response_model = smr_client.invoke_endpoint(
                EndpointName=endpoint_name,
                Body=json.dumps(
                {
                    "inputs": text_arrs,
                    "parameters": parameters,
                    "is_query" : False,
                    "instruction" :  None
                }
                ),
                ContentType="application/json",
            )
    
    json_str = response_model['Body'].read().decode('utf8')
    json_obj = json.loads(json_str)
    embeddings = json_obj["sentence_embeddings"]
    
    return embeddings

def batch_generator(generator, batch_size):
    while True:
        batch = list(itertools.islice(generator, batch_size))
        if not batch:
            break
        yield batch

def iterate_paragraph(file_content, object_key, doc_classify,smr_client, index_name, endpoint_name):
    json_arr = json.loads(file_content)
    doc_title = object_key
    text_splitter = RecursiveCharacterTextSplitter(        
        chunk_size = CHUNK_SIZE,
        chunk_overlap  = CHUNK_OVERLAP,
        length_function = len,
    )
    # def chunk_generator(json_arr):
    #     for idx, json_item in enumerate(json_arr):
    #         header = ""
    #         if len(json_item['heading']) > 0:
    #             header = json_item['heading'][0]['heading']

    #         paragraph_content = json_item['content']
    #         if len(paragraph_content) > 1024 or len(paragraph_content) < Sentence_Len_Threshold:
    #             continue

    #         yield (idx, paragraph_content, 'Paragraph', paragraph_content)

    #         sentences = re.split('[。？?.！!]', paragraph_content)
    #         for sent in (sent for sent in sentences if len(sent) > Sentence_Len_Threshold): 
    #             yield (idx, sent, 'Sentence', paragraph_content)

    def chunk_generator(json_arr):
        idx = 0
        texts = []
        for json_item in json_arr:
            header = ""
            if len(json_item['heading']) > 0:
                header = json_item['heading'][0]['heading']
            texts += text_splitter.split_text(f"{header}-{json_item['content']}")
            for paragraph_content in texts:
                idx += 1
                yield (idx, paragraph_content, 'Paragraph', paragraph_content)

    generator = chunk_generator(json_arr)
    batches = batch_generator(generator, batch_size=EMB_BATCH_SIZE)
    doc_author = get_filename_from_obj_key(object_key)

    for batch in batches:
        if batch is not None:
            emb_src_texts = [item[1] for item in batch]
            print("len of emb_src_texts :{}".format(len(emb_src_texts)))
            embeddings = get_embedding(smr_client, emb_src_texts, endpoint_name)
            for i, emb in enumerate(embeddings):
                document = { "publish_date": publish_date, "idx": batch[i][0], "doc" : batch[i][1], "doc_type" : batch[i][2], "content" : batch[i][3], "doc_title": doc_title,"doc_author":doc_author,"doc_classify":doc_classify, "doc_category": doc_title, "embedding" : emb}
                yield {"_index": index_name, "_source": document, "_id": hashlib.md5(str(document['doc']).encode('utf-8')).hexdigest()}

def iterate_pdf_json(file_content, object_key,doc_classify, smr_client, index_name, endpoint_name):
    json_arr = json.loads(file_content)
    doc_title = json_arr[0]['doc_title']

    def chunk_generator(json_arr):
        for idx, json_item in enumerate(json_arr):
            paragraph_content = None
            content = json_item['content']
            # print("----{}----".format(idx))
            # print(content)

            is_table = not isinstance(content, str)
            doc_category = 'table' if is_table else 'paragraph'
            if is_table:
                paragraph_content = "Table - {}\n{}\n\n{}".format(content['table'], json.dumps(content['data']), content['footer'])
            else:
                paragraph_content = "#{}\n{}".format(doc_title, content)
                if len(paragraph_content) > 1024 or len(paragraph_content) < Paragraph_Len_Threshold:
                    continue

            yield (idx, paragraph_content, 'Paragraph', paragraph_content, doc_category)

            if is_table:
                yield (idx, content['footer'], 'Sentence', content['footer'], doc_category)
            else:
                sentences = re.split('[。？?.！!]', paragraph_content)
                for sent in (sent for sent in sentences if len(sent) > Sentence_Len_Threshold): 
                    yield (idx, sent, 'Sentence', paragraph_content, doc_category)

    generator = chunk_generator(json_arr)
    batches = batch_generator(generator, batch_size=EMB_BATCH_SIZE)
    doc_author = get_filename_from_obj_key(object_key)
    try:
        for batch in batches:
            if batch is not None:
                emb_src_texts = [item[1] for item in batch]
                print("len of emb_src_texts :{}".format(len(emb_src_texts)))
                embeddings = get_embedding(smr_client, emb_src_texts, endpoint_name)
                for i, emb in enumerate(embeddings):
                    document = { "publish_date": publish_date, "idx": batch[i][0], "doc" : batch[i][1], "doc_type" : batch[i][2], "content" : batch[i][3], "doc_title": doc_title,"doc_author":doc_author, "doc_category": batch[i][4], "doc_classify":doc_classify,"embedding" : emb}
                    yield {"_index": index_name, "_source": document, "_id": hashlib.md5(str(document['doc']).encode('utf-8')).hexdigest()}
    except Exception as e:
        logging.exception(e)

def iterate_QA(file_content, object_key,doc_classify,smr_client, index_name, endpoint_name):
    json_content = json.loads(file_content)
    json_arr = json_content["qa_list"]
    doc_title = object_key
    doc_category = json_content["doc_category"]

    it = iter(json_arr)
    qa_batches = batch_generator(it, batch_size=EMB_BATCH_SIZE)
    doc_author = get_filename_from_obj_key(object_key)

    for idx, batch in enumerate(qa_batches):
        print(f"processing {idx*EMB_BATCH_SIZE} - {(idx+1)*EMB_BATCH_SIZE}")
        try:
            doc_type_list = [ item.get('doc_type', None) for item in batch ]
            doc_template = "Question: {}\nAnswer: {}"
            meta = [ { k:item[k] for k in item.keys() if k not in ['Question', 'Answer', 'Trigger', 'Content', 'Author', 'doc_type'] } for item in batch ]
            authors = [ item.get('Author')  for item in batch ]
            
            NoQA_embedding = set(doc_type_list)== {"NoQA"}
            if NoQA_embedding:
                contents = [ item['Content'] for item in batch ]
                triggers = [ item['Trigger'] for item in batch ]
                embeddings_trigger = get_embedding(smr_client, triggers, endpoint_name)
                for i in range(len(doc_type_list)):
                    document = { "publish_date": publish_date, "doc" : triggers[i], "idx": idx, "doc_type" : doc_type_list[i], "content" : contents[i], "doc_title": doc_title,"doc_author":authors[i] if authors[i] else doc_author, "doc_category": doc_category, "doc_meta": json.dumps(meta[i], ensure_ascii=False), "doc_classify":doc_classify,"embedding" : embeddings_trigger[i]}
                    yield {"_index": index_name, "_source": document, "_id": hashlib.md5(str(document).encode('utf-8')).hexdigest()}
            else:
                questions = [ item['Question'] for item in batch ]
                answers = [ item['Answer'] for item in batch ]
                contents = [ doc_template.format(item['Question'], item['Answer']) for item in batch ]

                embeddings_q = get_embedding(smr_client, questions, endpoint_name)
                embeddings_a = get_embedding(smr_client, answers, endpoint_name)
                for i in range(len(doc_type_list)):
                    document = { "publish_date": publish_date, "doc" : questions[i], "idx": idx, "doc_type" : "Question", "content" : contents[i], "doc_title": doc_title,"doc_author":authors[i] if authors[i] else doc_author, "doc_category": doc_category, "doc_meta": json.dumps(meta[i], ensure_ascii=False), "doc_classify":doc_classify,"embedding" : embeddings_q[i]}
                    yield {"_index": index_name, "_source": document, "_id": hashlib.md5(str(document).encode('utf-8')).hexdigest()}
                    document = { "publish_date": publish_date, "doc" : answers[i], "idx": idx,"doc_type" : "Paragraph", "content" : contents[i], "doc_title": doc_title,"doc_author":authors[i] if authors[i] else doc_author, "doc_category": doc_category, "doc_meta": json.dumps(meta[i], ensure_ascii=False), "doc_classify":doc_classify,"embedding" : embeddings_a[i]}
                    yield {"_index": index_name, "_source": document, "_id": hashlib.md5(str(document).encode('utf-8')).hexdigest()}
        except Exception as e:
            print(f"failed to process, {str(e)}")

def iterate_examples(file_content, object_key, smr_client, index_name, endpoint_name):
    json_obj = json.loads(file_content)

    api_schema = json_obj["api_schema"] if "api_schema" in json_obj.keys() else ""
    json_arr = json_obj["examples"]
    doc_title = object_key
    
    # print("api_schema:")
    # print(api_schema)
    # print("json_arr:")
    # print(json_arr)

    it = iter(json_arr)
    example_batches = batch_generator(it, batch_size=EMB_BATCH_SIZE)

    for idx, batch in enumerate(example_batches):
        queries = [ item['query'] for item in batch ]
        detections = [ json.dumps(item['detection'], ensure_ascii=False) for item in batch ]

        embeddings = get_embedding(smr_client, queries, endpoint_name)
        for i, query in enumerate(queries):
            # print("query:")
            # print(query)
            document = { "publish_date": publish_date, "detection" : detections[i], "query" : queries[i], "api_schema" : json.dumps(api_schema, ensure_ascii=False), "doc_title":doc_title, "embedding" : embeddings[i]}
            yield {"_index": index_name, "_source": document, "_id": hashlib.md5(str(document).encode('utf-8')).hexdigest()}
            
def link_header(semantic_snippets):
    heading_fonts_arr = [ item.metadata['heading_font'] for item in semantic_snippets ]
    heading_arr = [ item.metadata['heading'] for item in semantic_snippets ]

    def fontsize_mapping(heading_fonts_arr):
        heading_fonts_set = list(set(heading_fonts_arr))
        heading_fonts_set.sort(reverse=True)
        idxs = range(len(heading_fonts_set))
        font_idx_mapping = dict(zip(heading_fonts_set,idxs))
        return font_idx_mapping
        
    fontsize_dict = fontsize_mapping(heading_fonts_arr)

    snippet_arr = []
    for idx, snippet in enumerate(semantic_snippets):
        font_size = heading_fonts_arr[idx]
        heading_stack = []
        heading_info = {"font_size":heading_fonts_arr[idx], "heading":heading_arr[idx], "fontsize_idx" : fontsize_dict[font_size]}
        heading_stack.append(heading_info)
        for id in range(0,idx)[::-1]:
            if font_size < heading_fonts_arr[id]:
                font_size = heading_fonts_arr[id]
                heading_info = {"font_size":font_size, "heading":heading_arr[id], "fontsize_idx" : fontsize_dict[font_size]}
                heading_stack.append(heading_info)
            
        snippet_info = {
            "heading" : heading_stack,
            "content" : snippet.page_content
        }
        snippet_arr.append(snippet_info)
        
    json_arr = json.dumps(snippet_arr, ensure_ascii=False)
    return json_arr

def parse_pdf_to_json(file_content):
    soup = BeautifulSoup(file_content,'html.parser')
    content = soup.find_all('div')

    cur_fs = None
    cur_text = ''
    snippets = []   # first collect all snippets that have the same font size
    for c in content:
        sp = c.find('span')
        if not sp:
            continue
        st = sp.get('style')
        if not st:
            continue
        fs = re.findall('font-size:(\d+)px',st)
        if not fs:
            continue
        fs = int(fs[0])
        if not cur_fs:
            cur_fs = fs
        if fs == cur_fs:
            cur_text += c.text
        else:
            snippets.append((cur_text,cur_fs))
            cur_fs = fs
            cur_text = c.text
    snippets.append((cur_text,cur_fs))

    cur_idx = -1
    semantic_snippets = []
    # Assumption: headings have higher font size than their respective content
    for s in snippets:
        # if current snippet's font size > previous section's heading => it is a new heading
        if not semantic_snippets or s[1] > semantic_snippets[cur_idx].metadata['heading_font']:
            metadata={'heading':s[0], 'content_font': 0, 'heading_font': s[1]}
            #metadata.update(data.metadata)
            semantic_snippets.append(Document(page_content='',metadata=metadata))
            cur_idx += 1
            continue
        
        # if current snippet's font size <= previous section's content => content belongs to the same section (one can also create
        # a tree like structure for sub sections if needed but that may require some more thinking and may be data specific)
        if not semantic_snippets[cur_idx].metadata['content_font'] or s[1] <= semantic_snippets[cur_idx].metadata['content_font']:
            semantic_snippets[cur_idx].page_content += s[0]
            semantic_snippets[cur_idx].metadata['content_font'] = max(s[1], semantic_snippets[cur_idx].metadata['content_font'])
            continue
        
        # if current snippet's font size > previous section's content but less tha previous section's heading than also make a new 
        # section (e.g. title of a pdf will have the highest font size but we don't want it to subsume all sections)
        metadata={'heading':s[0], 'content_font': 0, 'heading_font': s[1]}
        #metadata.update(data.metadata)
        semantic_snippets.append(Document(page_content='',metadata=metadata))
        cur_idx += 1

    json_content = link_header(semantic_snippets)
    return json_content

def parse_faq_to_json(file_content):
    arr = file_content.split(QA_SEP)
    json_arr = []
    for item in arr:
        print(item)
        question, answer = item.strip().split("\n", 1)
        question = question.replace("Question: ", "")
        answer = answer.replace("Answer: ", "")
        obj = {
            "Question":question, "Answer":answer
        }
        json_arr.append(obj)

    qa_content = {
        "doc_title" : "",
        "doc_category" : "FAQ",
        "qa_list" : json_arr
    }
    
    json_content = json.dumps(qa_content, ensure_ascii=False)
    return json_content
    
def parse_txt_to_json(file_content):
    text_splitter = RecursiveCharacterTextSplitter( 
        chunk_size = arg_chunk_size,
        chunk_overlap  = 0,
    )
    
    results = []
    chunks = text_splitter.create_documents([ file_content ] )
    for chunk in chunks:
        snippet_info = {
            "heading" : [],
            "content" : chunk.page_content
        }
        results.append(snippet_info)

    json_content = json.dumps(results, ensure_ascii=False)
    return json_content

def parse_md_to_json(file_content):
    md_splitter = MarkdownTextSplitter( 
        chunk_size = arg_chunk_size,
        chunk_overlap  = 0,
    )
    
    results = []
    chunks = md_splitter.create_documents([ file_content ] )
    for chunk in chunks:
        snippet_info = {
            "heading" : [],
            "content" : chunk.page_content
        }
        results.append(snippet_info)

    json_content = json.dumps(results, ensure_ascii=False)
    return json_content

def parse_example_to_json(file_content):
    arr = file_content.split(EXAMPLE_SEP)
    json_arr = []

    for item in arr:
        elements = item.strip().split("\n")
        print("elements:")
        print(elements)
        obj = { element.split(":")[0] : element.split(":")[1] for element in elements }
        json_arr.append(obj)

    qa_content = {
        "example_list" : json_arr
    }
    
    json_content = json.dumps(qa_content, ensure_ascii=False)
    return json_content

def parse_html_to_json(html_docs):
    text_splitter = RecursiveCharacterTextSplitter( 
        chunk_size = arg_chunk_size,
        chunk_overlap  = 0,
    )

    results = []
    chunks = text_splitter.create_documents([ doc.page_content for doc in html_docs ] )
    for chunk in chunks:
        snippet_info = {
            "heading" : [],
            "content" : chunk.page_content
        }
        results.append(snippet_info)
    json_content = json.dumps(results, ensure_ascii=False)
    return json_content


def parse_xlsx_to_json(file_content):
    import pandas as pd
    import io

    df = pd.read_excel(io.BytesIO(file_content))
    df.fillna(value='', inplace=True)

    json_arr = df.to_dict(orient='records')
    qa_content = {
        "doc_title" : "",
        "doc_category" : "FAQ",
        "qa_list" : json_arr
    }

    json_content = json.dumps(qa_content, ensure_ascii=False)
    return json_content

## parse faq in csv format. Question    Answer
def parse_csv_to_json(file_content):
    import csv
    csv_data = file_content.splitlines()
    reader = csv.reader(csv_data)
    header = next(reader)  # Skip the header row
    json_arr = []
    for item in reader:
        if len(item) >=3:
            question, answer,author = item[0],item[1],item[2]
            question = question.replace("Question: ", "")
            answer = answer.replace("Answer: ", "")
            obj = {
                "Question":question, "Answer":answer,"Author":author
            }
            json_arr.append(obj)
        elif len(item) ==2:
            question, answer = item[0],item[1]
            question = question.replace("Question: ", "")
            answer = answer.replace("Answer: ", "")
            obj = {
                "Question":question, "Answer":answer
            }
            json_arr.append(obj)
        else:
            raise ('csv file must have two columns at least')
            

    qa_content = {
        "doc_title" : "",
        "doc_category" : "FAQ",
        "qa_list" : json_arr
    }
    
    json_content = json.dumps(qa_content, ensure_ascii=False)
    return json_content

def parse_docx_to_json(file_content):
    import docx
    doc = docx.Document(file_content)
    text = ''
    for para in doc.paragraphs:
        text += '\n'+para.text
    text_splitter = RecursiveCharacterTextSplitter(        
        chunk_size = CHUNK_SIZE,
        chunk_overlap  = CHUNK_OVERLAP,
        length_function = token_length_function,
    )
    results = []
    chunks = text_splitter.create_documents([ text ] )
    for chunk in chunks:
        snippet_info = {
            "heading" : [],
            "content" : chunk.page_content
        }
        results.append(snippet_info)

    json_content = json.dumps(results, ensure_ascii=False)
    return json_content


def load_content_json_from_s3(bucket, object_key, content_type, credentials):
    if content_type in ['pdf','docx']:
        pdf_path=os.path.basename(object_key)
        s3_client=boto3.client('s3', region_name=REGION)
        s3_client.download_file(Bucket=bucket, Key=object_key, Filename=pdf_path)
        if content_type == 'pdf':
            loader = PDFMinerPDFasHTMLLoader(pdf_path)
            file_content = loader.load()[0].page_content
            json_content = parse_pdf_to_json(file_content)
            return json_content
        elif content_type == 'docx':
            json_content = parse_docx_to_json(pdf_path)
            return json_content
    else:
        s3 = boto3.resource('s3')
        obj = s3.Object(bucket,object_key)
        if content_type != 'xlsx':
            file_content = obj.get()['Body'].read().decode('utf-8', errors='ignore').strip()
        else:
            # binrary mode
            file_content = obj.get()['Body'].read()
        try:
            if content_type == 'faq':
                json_content = parse_faq_to_json(file_content)
            elif content_type == 'md':
                json_content = parse_md_to_json(file_content)
            elif content_type =='txt':
                json_content = parse_txt_to_json(file_content)
            elif content_type =='json':
                json_content = file_content
            elif content_type == 'pdf.json':
                json_content = file_content
            elif content_type == 'example':
                json_content = file_content
            elif content_type in ['wiki','blog']:
                json_content = json.loads(file_content)
            elif content_type == 'csv':
                json_content = parse_csv_to_json(file_content)
            elif content_type == 'xlsx':
                json_content = parse_xlsx_to_json(file_content)
            else:
                print(f"unsupport content type...{content_type}")
                raise RuntimeError(f"unsupport content type...{content_type}")
        except Exception as e:
            raise RuntimeError(f"Exception ...{str(e)}")
        
        return json_content

def iterate_paragraph_blog(content_json, object_key,doc_classify,smr_client, index_name, endpoint_name):
    doc_title = object_key
    text_splitter = RecursiveCharacterTextSplitter(        
        chunk_size = CHUNK_SIZE,
        chunk_overlap  = CHUNK_OVERLAP,
        length_function = token_length_function,
    )
    def chunk_generator(json_arr):
        for blog in json_arr:
            idx = 0
            doc_title = blog['doc_title']
            content_summary = blog['content_summary']
            texts = text_splitter.split_text(f"{content_summary}")
            for parag in blog['content']:
                texts += text_splitter.split_text(f"{parag['title']}-{parag['content']}")
            for paragraph_content in texts:
                idx += 1
                yield (idx, paragraph_content, 'Paragraph', paragraph_content,doc_title)
                ## add embedding for sentence
                # 实测效果并不好，造成召回内容干扰
                # sentences = re.split('[。？?.！!]', paragraph_content)
                # for sent in (sent for sent in sentences if len(sent) > Sentence_Len_Threshold): 
                #     yield (idx, sent, 'Sentence', paragraph_content,doc_title)

    generator = chunk_generator(content_json)
    batches = batch_generator(generator, batch_size=EMB_BATCH_SIZE)
    doc_author = get_filename_from_obj_key(object_key)
    for batch in batches:
        if batch is not None:
            emb_src_texts = [item[1] for item in batch] ##对content向量化
            print("len of emb_src_texts :{}".format(len(emb_src_texts)))
            embeddings = get_embedding(smr_client, emb_src_texts, endpoint_name)
            for i, emb in enumerate(embeddings):
                document = { "publish_date": publish_date, "idx": batch[i][0], "doc" : batch[i][1], "doc_type" : batch[i][2], "content" : batch[i][3], "doc_title": doc_title, "doc_author":doc_author,"doc_category": batch[i][4], "doc_classify":doc_classify,"embedding" : emb}
                yield {"_index": index_name, "_source": document, "_id": hashlib.md5(str(document['doc']).encode('utf-8')).hexdigest()}
                

def iterate_paragraph_wiki(content_json, object_key,doc_classify,smr_client, index_name, endpoint_name):
    doc_title = object_key
    text_splitter = RecursiveCharacterTextSplitter(        
        chunk_size = CHUNK_SIZE,
        chunk_overlap  = CHUNK_OVERLAP,
        length_function = token_length_function,
    )
    def chunk_generator(json_arr):
        for page in json_arr:
            idx = 0
            for url,p in page.items():
                texts = []
                if len(p):
                    print(f'page:{url}, size: {len(p[0])}')
                    texts = text_splitter.split_text(p[0])
                for paragraph_content in texts:
                    idx += 1
                    yield (idx, paragraph_content, 'Paragraph', paragraph_content,url)
                    ## add embedding for sentence
                    # 实测效果并不好，造成召回内容干扰
                    # sentences = re.split('[。？?.！!]', paragraph_content)
                    # for sent in (sent for sent in sentences if len(sent) > Sentence_Len_Threshold): 
                    #     yield (idx, sent, 'Sentence', paragraph_content,doc_title,url)

    generator = chunk_generator(content_json)
    batches = batch_generator(generator, batch_size=EMB_BATCH_SIZE)
    doc_author = get_filename_from_obj_key(object_key)
    for batch in batches:
        if batch is not None:
            emb_src_texts = [item[1] for item in batch] ##对content向量化
            print("len of emb_src_texts :{}".format(len(emb_src_texts)))
            embeddings = get_embedding(smr_client, emb_src_texts, endpoint_name)
            for i, emb in enumerate(embeddings):
                document = { "publish_date": publish_date, "idx": batch[i][0], "doc" : batch[i][1], "doc_type" : batch[i][2], "content" : batch[i][3], "doc_title": doc_title, "doc_author":doc_author,"doc_category": batch[i][4], "doc_classify":doc_classify,"embedding" : emb}
                yield {"_index": index_name, "_source": document, "_id": hashlib.md5(str(document['doc']).encode('utf-8')).hexdigest()}



def put_idx_to_ddb(filename,company,username,index_name,embedding_model,category,createtime):
    try:
        dynamodb.put_item(
            Item={
                'filename':{
                    'S':filename,
                },
                'username':{
                    'S':username,
                },
                'company':{
                    'S':company,
                },
                'index_name':{
                    'S':index_name,
                },
                'embedding_model':{
                    'S':embedding_model,
                },
                'category':{
                    'S':category,
                },
                'createtime':{
                    'S':createtime,
                }
            },
            TableName = DOC_INDEX_TABLE,
        )
        print(f"Put filename:{filename} with embedding:{embedding_model} index_name:{index_name} by user:{username} to ddb success")
        return True
    except Exception as e:
        print(f"Not found filename:{filename} with embedding:{embedding_model} index_name:{index_name} to ddb: {str(e)}")
        return False 


def query_idx_from_ddb(filename,username,embedding_model):
    try:
        response = dynamodb.query(
            TableName=DOC_INDEX_TABLE,
            ExpressionAttributeValues={
                ':v1': {
                    'S': filename,
                },
                ':v2': {
                    'S': username,
                },
                ':v3': {
                    'S': embedding_model,
                }
            },
            KeyConditionExpression='filename = :v1 and username = :v2',
            ExpressionAttributeNames={"#e":"embedding_model"},
            FilterExpression='#e = :v3',
            ProjectionExpression='index_name'
        )
        if len(response['Items']):
            index_name = response['Items'][0]['index_name']['S'] 
        else:
            index_name = ''
        print (f"query filename:{filename} with embedding:{embedding_model} index_name:{index_name} from ddb")
        return index_name
    
    except Exception as e:
        print(f"Not found filename:{filename}, username:{username}, embedding_model:{embedding_model} index from ddb")
        return ''

    
def WriteVecIndexToAOS(bucket, object_key, content_type, doc_classify,smr_client, aos_endpoint=AOS_ENDPOINT, region=REGION, index_name=INDEX_NAME):
    credentials = boto3.Session().get_credentials()
    auth = AWSV4SignerAuth(credentials, region)
    # auth = ('xxxx', 'yyyy') master user/pwd
    # auth = (aos_master, aos_pwd)
    try:
        file_content = load_content_json_from_s3(bucket, object_key, content_type, credentials)
        # print("file_content:")
        # print(file_content)

        client = OpenSearch(
            hosts = [{'host': aos_endpoint, 'port': 443}],
            http_auth = auth,
            use_ssl = True,
            verify_certs = True,
            connection_class = RequestsHttpConnection,
            timeout = 60, # 默认超时时间是10 秒，
            max_retries=5, # 重试次数
            retry_on_timeout=True
        )

        print("---------flag------")
        gen_aos_record_func = None
        if content_type in ["faq","csv","xlsx"]:
            gen_aos_record_func = iterate_QA(file_content, object_key,doc_classify,smr_client, index_name, EMB_MODEL_ENDPOINT)
        elif content_type in ['txt', 'pdf', 'json','docx', 'md']:
            gen_aos_record_func = iterate_paragraph(file_content,object_key, doc_classify,smr_client, index_name, EMB_MODEL_ENDPOINT)
        elif content_type in [ 'pdf.json' ]:
            gen_aos_record_func = iterate_pdf_json(file_content,object_key, doc_classify,smr_client, index_name, EMB_MODEL_ENDPOINT)
        elif content_type in ['example']:
            gen_aos_record_func = iterate_examples(file_content,object_key,smr_client, index_name, EMB_MODEL_ENDPOINT)
        elif content_type in ['wiki']:
            gen_aos_record_func = iterate_paragraph_wiki(file_content,object_key,doc_classify, smr_client, index_name, EMB_MODEL_ENDPOINT)
        elif content_type in ['blog']:
            gen_aos_record_func = iterate_paragraph_blog(file_content,object_key,doc_classify, smr_client, index_name, EMB_MODEL_ENDPOINT)
        else:
            raise RuntimeError('No Such Content type supported') 

        # chunk_size 为文档数 默认值为500
        # max_chunk_bytes 为写入的最大字节数，默认100M过大，可以改成10-15M
        # max_retries 重试次数
        # initial_backoff 为第一次重试时sleep的秒数，再次重试会翻倍
        response = helpers.bulk(client, gen_aos_record_func, max_retries=3, initial_backoff=200, max_backoff=801, max_chunk_bytes=10 * 1024 * 1024)#, chunk_size=10000, request_timeout=60000) 
        return response
    except Exception as e:
        print(f"There was an error when ingest:{object_key} to aos cluster, Exception: {str(e)}")
        return None   

def process_s3_uploaded_file(bucket, object_key):
    print("********** object_key : " + object_key)
    #if want to use different aos index, the object_key format should be: ai-content/company/username/filename
            
    content_type = None
    if object_key.endswith(".faq"):
        print("********** pre-processing faq file")
        content_type = 'faq'
    elif object_key.endswith(".txt"):
        print("********** pre-processing text file")
        content_type = 'txt'
    elif re.search(r'.wiki(\(\d+\))*.json',object_key):
        print("********** pre-processing wiki file")
        content_type = 'wiki'
    elif re.search(r'.blog(\(\d+\))*.json',object_key):
        print("********** pre-processing blog file")
        content_type = 'blog'
    elif re.search(r'.pdf(\(\d+\))*.json',object_key):
        print("********** pre-processing pdf.json file")
        content_type = 'pdf.json'
    elif object_key.endswith(".pdf"):
        print("********** pre-processing pdf file")
        content_type = 'pdf'
    elif object_key.endswith(".json"):
        print("********** pre-processing json file")
        content_type = 'json'
    elif object_key.endswith(".example"):
        print("********** pre-processing example file")
        content_type = 'example'
    elif object_key.endswith(".csv"):
        print("********** pre-processing csv file")
        content_type = 'csv'
    elif object_key.endswith(".docx"):
        print("********** pre-processing docx file")
        content_type = 'docx'
    elif object_key.endswith(".xlsx"):
        print("********** pre-processing xlsx file")
        content_type = 'xlsx'
    elif object_key.endswith(".md"):
        print("********** pre-processing md file")
        content_type = 'md'
    else:
        print(f"{object_key} - unsupport content type...(pdf, faq, txt, csv, xlsx, pdf.json, md are supported.)")
        return
        
    username = get_filename_from_obj_key(object_key)
    
    s3 = boto3.resource('s3')
    obj = s3.Object(bucket,object_key)
    metadata = obj.metadata
    

    doc_classify = unquote(metadata.get('category','') if metadata else '')
    # COMPANY should passed by args
    if  content_type == 'example':
        index_name = f"chatbot-example-index-{COMPANY}"
    else:
        index_name = f"{INDEX_NAME}-{COMPANY}"


        
    print(metadata)
    print(f'use index :{index_name}')
    #check if it is already built
    idx_name = query_idx_from_ddb(object_key,username,EMB_MODEL_ENDPOINT)
    if len(idx_name) > 0:
        print("doc file already exists")
        return
    

    response = WriteVecIndexToAOS(bucket, object_key, content_type,doc_classify, smr_client, index_name=index_name)
    print("response:")
    print(response)
    print("ingest {} chunk to AOS".format(response[0]))
    timestamp_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    put_idx_to_ddb(filename=object_key,
                    company=COMPANY,
                    username=username,
                        index_name=index_name,
                            embedding_model=EMB_MODEL_ENDPOINT,
                            category=doc_classify,
                            createtime=timestamp_str)

##如果是从chatbot上传，则是ai-content/username/filename
def get_filename_from_obj_key(object_key):
    paths = object_key.split('/')
    return paths[1] if len(paths) > 2 else 's3_upload'



for s3_key in object_key.split(','):
    s3_key = urllib.parse.unquote(s3_key) ##In case Chinese filename
    s3_key = s3_key.replace('+',' ') ##replace the '+' with space. ps:if the original file name contains space, then s3 notification will replace it with '+'.
    print("processing {}".format(s3_key))
    process_s3_uploaded_file(bucket, s3_key)