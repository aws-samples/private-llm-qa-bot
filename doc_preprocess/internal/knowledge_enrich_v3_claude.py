import json
import os
import logging
from collections import Counter
import xml.etree.ElementTree as ET
from langchain.document_loaders import PDFMinerPDFasHTMLLoader
from langchain.prompts import PromptTemplate
from typing import Any, Dict, List, Union,Mapping, Optional, TypeVar, Union
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.chains import LLMChain
from langchain.llms.bedrock import Bedrock
from botocore.exceptions import ClientError
import boto3
import math
import csv
import re
import argparse
from anthropic_bedrock import AnthropicBedrock
import docx
from bs4 import BeautifulSoup
from langchain.docstore.document import Document


num_size = 200
bedrock = AnthropicBedrock(
    aws_region="us-west-2",

)


boto3_bedrock = boto3.client(
        service_name="bedrock-runtime",
        region_name='us-west-2'
    )


prompt_template = \
"""
Please summarize below content in <content></content> to generate {num} standalone questions and correspondant answers, 
response in xml format,for example:
<faqs>
<faq>
<question>xxx</question>
<answer>xxx</answer>
<faq>
<question>xxx</question>
<answer>xxx</answer>
</faq>
<faq>
<question>xxx</question>
<answer>xxx</answer>
</faq>
</faqs>

<content>
{content}
</content>
"""

prompt_template_summarize = \
"""
Here is a question and its corresponding answer: 
<question>
{question}
</quenstion>
<answer>
{answer}
</answer>
I'd like you to rewrite the answer paragraph using the following instructions: "less than 100 words, as acurate as the original answer".
Please put your rewrite in <rewrite></rewrite> tags.
"""

prompt_template_rewrite_qa = \
"""
Here is a input question: 
<question>
{question}
</quenstion>
I'd like you to rewrite the question using the following instructions: 
"generate less than 50 words, be more colloquial and keep the same meaning as the original question,
if the original quesiton is too long or divegent, please split in to several questions, and one question per line
".
Please put your rewrite in <rewrite></rewrite> tags. please response in the same language as the input question.
"""


def parse_xml_faq(content):
    pattern = r"<faqs>(.*?)</faqs>"
    matches = re.search(pattern, content,re.DOTALL)
    q_and_a = []
    if matches:
        content = matches.group()
        print(content)
        try:
            tree = ET.fromstring(content.strip())
            for faq in tree.iter('faq'):
                q_and_a.append([faq.find("question").text, faq.find("answer").text,'synthetic'])
        except Exception as e:
            print(str(e))
    return q_and_a

def parse_faq(faq_text):
    """
    解析 FAQ 文本并返回表格数据。

    Args:
    faq_text: FAQ 文本。

    Returns:
    表格数据。
    """

    # 匹配问题和答案的正则表达式
    pattern = r"(Q\d+):(.*)\n*(A\d+):(.\n*)"

    # 使用正则表达式匹配问题和答案
    matches = re.findall(pattern, faq_text)

    # 创建表格数据
    table_data = []
    for match in matches:
        table_data.append([match[1].strip(), match[3].strip(),'synthetic'])
    return table_data

def write_faq_pairs_to_csv(faq_pairs:list, csv_file:str,headers:list):
    with open(csv_file, 'w', newline='',encoding='utf-8-sig') as file:
        writer = csv.writer(file)
        writer.writerow(headers)
        writer.writerows(faq_pairs)



def num_tokens_from_string(string: str) -> int:
    """Returns the number of tokens in a text string."""
    num_tokens = bedrock.count_tokens(string)  # 
    # print(f"num_tokens:{num_tokens}")
    return num_tokens

def parse_df_text(file_name:str):
    loader = PDFMinerPDFasHTMLLoader(file_name)
    data = loader.load()[0]
    soup = BeautifulSoup(data.page_content,'html.parser')
    content = soup.find_all('div')
    cur_fs = None
    cur_text = ''
    snippets = []   # first collect all snippets that have the same font size
    for c in content:
        sp = c.find('span')
        if not sp:
            continue
        st = sp.get('style')
        if not st:
            continue
        fs = re.findall('font-size:(\d+)px',st)
        if not fs:
            continue
        fs = int(fs[0])
        if not cur_fs:
            cur_fs = fs
        if fs == cur_fs:
            cur_text += c.text
        else:
            snippets.append((cur_text,cur_fs))
            cur_fs = fs
            cur_text = c.text
    snippets.append((cur_text,cur_fs))
    cur_idx = -1
    semantic_snippets = []
    # Assumption: headings have higher font size than their respective content
    for s in snippets:
        # if current snippet's font size > previous section's heading => it is a new heading
        if not semantic_snippets or s[1] > semantic_snippets[cur_idx].metadata['heading_font']:
            metadata={'heading':s[0], 'content_font': 0, 'heading_font': s[1]}
            metadata.update(data.metadata)
            semantic_snippets.append(Document(page_content='',metadata=metadata))
            cur_idx += 1
            continue

        # if current snippet's font size <= previous section's content => content belongs to the same section (one can also create
        # a tree like structure for sub sections if needed but that may require some more thinking and may be data specific)
        if not semantic_snippets[cur_idx].metadata['content_font'] or s[1] <= semantic_snippets[cur_idx].metadata['content_font']:
            semantic_snippets[cur_idx].page_content += s[0]
            semantic_snippets[cur_idx].metadata['content_font'] = max(s[1], semantic_snippets[cur_idx].metadata['content_font'])
            continue

        # if current snippet's font size > previous section's content but less than previous section's heading than also make a new
        # section (e.g. title of a PDF will have the highest font size but we don't want it to subsume all sections)
        metadata={'heading':s[0], 'content_font': 0, 'heading_font': s[1]}
        metadata.update(data.metadata)
        semantic_snippets.append(Document(page_content='',metadata=metadata))
        cur_idx += 1
    return semantic_snippets
    


def generate(file_name,llmchain:LLMChain,file_type:str):
    faq_table = []
    if file_type == 'csv':
        with open(file_name, 'r' ) as file:
            csv_data = file.readlines()
            reader = csv.reader(csv_data)
            header = next(reader)  # Skip the header row
            for item in reader:
                question, answer = item[0],item[1]
                content = f"Question:{question}\nAnswer:{answer}"
                faq_pairs = synth_faq(content,llmchain)
                faq_table += faq_pairs if faq_pairs else []
                # break #for test 
    elif file_type in ['docx','pdf'] :
        text = ''
        if file_type == 'docx':
            doc = docx.Document(file_name)
            for para in doc.paragraphs:
                text += '\n'+para.text
        elif file_type == 'pdf':
            docs = parse_df_text(file_name)
            for i,doc in enumerate(docs):
                text += doc.page_content
                
        text_splitter = RecursiveCharacterTextSplitter(        
            chunk_size = 2000,
            chunk_overlap  = 50,
            length_function = num_tokens_from_string,
        )
        texts = text_splitter.split_text(text)
        print(f'-----total-chunks:{len(texts)}-----')
        for i,content in enumerate(texts):
            print(f'-----chunk:{i}--tokens:{num_tokens_from_string(content)}')
            print(content)
            faq_pairs = synth_faq(content,llmchain)
            faq_table += faq_pairs if faq_pairs else []    
        
    return faq_table



def summarize_qa(file_name:str,file_type:str):
    if file_type != 'csv':
        raise ('Only support csv')
    parameters = {
        "max_tokens_to_sample": 500,
        "stop_sequences": ["</output>"],
        "temperature":0.1,
        "top_p":1
    }
    llm = Bedrock(model_id='anthropic.claude-v2:1', client=boto3_bedrock, model_kwargs=parameters)

    PROMPT = PromptTemplate(
            template=prompt_template_summarize, 
            input_variables=['question','answer']
    )

    llmchain = LLMChain(llm=llm, verbose=False, prompt=PROMPT)
    
    def parse_rewrite(text:str):
        pattern = r"<rewrite>(.*?)</rewrite>"
        matches = re.search(pattern, text,re.DOTALL)
        return matches.group(1) if matches else ''

    faq_table = []
    with open(file_name, 'r' ) as file:
        csv_data = file.readlines()
        reader = csv.reader(csv_data)
        header = next(reader)  # Skip the header row
        for item in reader:
            question, answer,author = item[0],item[1],item[2]
            num_tokens = num_tokens_from_string(answer)
            if num_tokens > 200:
                print(f'num_tokens:{num_tokens}')
                response = llmchain.run({'question':question,'answer':answer})
                sum_a = parse_rewrite(response)
                faq_table.append ([question,answer,sum_a,author])
    return faq_table
                
def rewrite_qa(file_name:str,file_type:str):
    if file_type != 'csv':
        raise ('Only support csv')
    parameters = {
        "max_tokens_to_sample": 500,
        "stop_sequences": ["</output>"],
        "temperature":0.1,
        "top_p":1
    }
    llm = Bedrock(model_id='anthropic.claude-v2:1', client=boto3_bedrock, model_kwargs=parameters)

    PROMPT = PromptTemplate(
            template=prompt_template_rewrite_qa, 
            input_variables=['question']
    )

    llmchain = LLMChain(llm=llm, verbose=False, prompt=PROMPT)
    
    def parse_rewrite(text:str):
        pattern = r"<rewrite>(.*?)</rewrite>"
        matches = re.search(pattern, text,re.DOTALL)
        return matches.group(1) if matches else ''
    
    faq_table = []
    with open(file_name, 'r' ) as file:
        csv_data = file.readlines()
        reader = csv.reader(csv_data)
        header = next(reader)  # Skip the header row
        idx = 0
        for item in reader:
            question, answer = item[0],item[1]
            response = llmchain.run({'question':question})
            # print(response)
            new_q = parse_rewrite(response).strip()
            qs = new_q.split('\n')
            for q in qs:
                print('new:',q)
                faq_table.append ([q,answer,question,idx])
            idx += 1
    return faq_table
            
            
    
    


def synth_faq(content:str,llmchain:LLMChain):
    global num_size
    num_tokens = num_tokens_from_string(content)
    answer = llmchain.run({"content":content,"num":math.ceil(num_tokens/num_size)})
    # print(answer)
    return parse_xml_faq(answer)

def main():
    global num_size
    parser = argparse.ArgumentParser()
    parser.add_argument("--file", type=str,)
    parser.add_argument("--size", type=int, default=200)
    parser.add_argument("--summarize", type=bool, default=False)
    parser.add_argument("--rewrite", type=bool, default=False)
    args = parser.parse_args()
    num_size = args.size
    parameters = {
        "max_tokens_to_sample": 4000,
        "stop_sequences": ["</output>"],
        "temperature":0.1,
        "top_p":1
    }

    llm = Bedrock(model_id='anthropic.claude-v2:1', client=boto3_bedrock, model_kwargs=parameters)

    PROMPT = PromptTemplate(
            template=prompt_template, 
            input_variables=['num','content']
    )

    llmchain = LLMChain(llm=llm, verbose=False, prompt=PROMPT)

    input_file = args.file
    file_type = input_file.split(".")[-1]
    if args.summarize:
        new_faq = summarize_qa(input_file,file_type)
        outout_file = input_file.replace('.csv',"")+'_sum_answer.csv'
        write_faq_pairs_to_csv(new_faq,outout_file,headers = ['Question', 'Answer','Sum_Answer','Author'])
        print(f'generated:{outout_file}')
    
    elif args.rewrite:
        new_faq = rewrite_qa(input_file,file_type)
        outout_file = input_file.replace('.csv',"")+'_new_question.csv'
        write_faq_pairs_to_csv(new_faq,outout_file,headers = ['New_Question', 'Answer','Question','Index'])
        print(f'generated:{outout_file}')
        
    else:
        new_faq = generate(input_file,llmchain,file_type)
        outout_file = input_file.replace('.csv',"")+'_synth.csv'
        write_faq_pairs_to_csv(new_faq,outout_file,headers = ['Question', 'Answer','Author'])
        print(f'generated:{outout_file}')
    
if __name__ == '__main__':
    main()
    text = \
"""
here is :
<faqs>
<faq>
<question>How does DynamoDB offer better performance compared to Cosmos DB?</question>
<answer>DynamoDB offers up to 10X better performance on reads and 30% faster performance on writes than Cosmos DB. DynamoDB delivers single-digit millisecond latencies while Cosmos DB offers 10ms latency on reads and 15ms on writes at the 99th percentile.</answer>
</faq>
<faq>
<question>Why is Amazon Neptune better for graph use cases compared to Cosmos DB?</question>
<answer>Neptune has better support for common graph models like property graphs and RDF, optimized high performance architecture, and bulk import capabilities. Cosmos DB lacks RDF support, has degraded multi-hop query performance, and no bulk import.</answer>
</faq>
<faq>

<question>How does Amazon RDS compare to Azure Database for MySQL and PostgreSQL?</question>

<answer>RDS provides more database engine choices, optimized engines with Aurora, better enterprise capabilities like automated Multi-AZ replication, and greater configurability and scalability.</answer>

</faq>

<faq>

<question>What AWS service provides database migration capabilities comparable to Azure?</question>

<answer>AWS Database Migration Service enables streaming data replication and continuous migration of databases like Oracle, SQL Server, and MongoDB into AWS with minimal downtime.</answer>

</faq>

</faqs>
"""
    # print(parse_xml_faq(text))
#     def parse_rewrite(text:str):
#         pattern = r"<rewrite>(.*?)</rewrite>"
#         matches = re.search(pattern, text,re.DOTALL)
#         return matches.group(1) if matches else ''
#     text = """<rewrite>
# AWS offers over 15 purpose-built databases to support diverse workloads, allowing customers to optimize for scale, performance and cost. AWS provides multiple options for running Oracle workloads including the fully-managed Amazon RDS for Oracle, launched over 11 years ago, and Amazon RDS Custom for Oracle to manage legacy applications. Customers have successfully migrated Oracle applications like E-Business Suite and JD Edwards to AWS. Resources are available to assist with migrations of Oracle Applications and Exadata workloads.  
# </rewrite>"""
#     print(parse_rewrite(text))
    

 
